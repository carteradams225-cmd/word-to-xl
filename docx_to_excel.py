"""
docx_to_excel.py

Parses a Word document (table-based structure) and APPENDS new records
to an existing Excel master table. Each initiative gets one row, with
its update text placed in the appropriate section column.

Columns: Date | Division | Initiative | Progress | Plans | Problems

Dedup key: Division + Initiative (same initiative on same date is skipped).
If an initiative appears under multiple sections, each section's text
goes into the matching column on the same row.

Requirements (install once):
    pip install python-docx openpyxl

Usage:
    python docx_to_excel.py input.docx output.xlsx
"""

import sys
import os
from datetime import date
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

SECTION_NAMES = {"progress", "plans", "problems"}
HEADERS = ["Date", "Division", "Initiative", "Progress", "Plans", "Problems"]

# Column index (1-based) for each section
SECTION_COL = {"Progress": 4, "Plans": 5, "Problems": 6}


# ── Parsing helpers ───────────────────────────────────────────────────────────

def get_indent_level(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return 0
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return 0
    ilvl = numPr.find(qn("w:ilvl"))
    return int(ilvl.get(qn("w:val"), 0)) if ilvl is not None else 0


def is_list_paragraph(para):
    if "list" in para.style.name.lower() or "bullet" in para.style.name.lower():
        return True
    pPr = para._p.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def extract_bold_and_rest(para):
    bold_parts, rest_parts, switched = [], [], False
    for run in para.runs:
        if not switched and run.bold and run.text.strip():
            bold_parts.append(run.text)
        else:
            switched = True
            rest_parts.append(run.text)
    bold = "".join(bold_parts).strip().rstrip(":")
    rest = "".join(rest_parts).strip().lstrip(": ")
    return bold, rest


def is_division_header(para):
    text = para.text.strip()
    if not text or len(text) > 120:
        return False
    if text.lower() in SECTION_NAMES:
        return False
    if is_list_paragraph(para):
        return False
    return True


def parse_cell_paragraphs(cell_paragraphs, section, results):
    current_division = "Unknown"
    last_top_level = None
    found_division = False

    for para in cell_paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.lower() in SECTION_NAMES:
            continue

        if not found_division and not is_list_paragraph(para):
            current_division = text
            found_division = True
            continue

        if is_division_header(para) and not is_list_paragraph(para):
            current_division = text
            last_top_level = None
            continue

        if is_list_paragraph(para):
            indent = get_indent_level(para)
            bold, rest = extract_bold_and_rest(para)

            if indent == 0:
                entry = {
                    "section":    section,
                    "division":   current_division,
                    "initiative": bold if bold else text,
                    "update":     rest if bold else "",
                    "notes":      [],
                }
                results.append(entry)
                last_top_level = entry
            else:
                if last_top_level is not None:
                    note = (bold + (": " + rest if rest else "")) if bold else text
                    last_top_level["notes"].append(note)
        else:
            if found_division:
                bold, rest = extract_bold_and_rest(para)
                entry = {
                    "section":    section,
                    "division":   current_division,
                    "initiative": bold if bold else text,
                    "update":     rest if bold else "",
                    "notes":      [],
                }
                results.append(entry)
                last_top_level = entry


def parse_document(path):
    doc = Document(path)
    raw_results = []
    current_section = "Unknown"

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for cell in cells:
                if cell.text.strip().lower() in SECTION_NAMES:
                    current_section = cell.text.strip().capitalize()
                    break
            for cell in cells:
                cell_text = cell.text.strip()
                if cell_text.lower() in SECTION_NAMES or not cell_text:
                    continue
                parse_cell_paragraphs(cell.paragraphs, current_section, raw_results)

    # Flatten notes and group by (division, initiative) so each initiative = one row
    # with up to three section columns filled
    grouped = {}  # key: (division, initiative) -> {progress, plans, problems}
    order = []    # preserve first-seen order

    for e in raw_results:
        update = e["update"]
        if e["notes"]:
            notes_str = " | ".join(e["notes"])
            update = (update + " [Notes: " + notes_str + "]") if update else "[Notes: " + notes_str + "]"

        key = (e["division"], e["initiative"])
        if key not in grouped:
            grouped[key] = {"Progress": "", "Plans": "", "Problems": ""}
            order.append(key)

        sec = e["section"].capitalize()
        if sec in grouped[key]:
            grouped[key][sec] = update

    result = []
    for (division, initiative) in order:
        cols = grouped[(division, initiative)]
        result.append({
            "division":   division,
            "initiative": initiative,
            "Progress":   cols["Progress"],
            "Plans":      cols["Plans"],
            "Problems":   cols["Problems"],
        })

    return result


# ── Excel helpers ─────────────────────────────────────────────────────────────

def header_border():
    s = Side(style="medium", color="1F4E79")
    return Border(left=s, right=s, top=s, bottom=s)


def cell_border():
    s = Side(style="thin", color="BFBFBF")
    return Border(left=s, right=s, top=s, bottom=s)


def style_header_row(ws):
    # Base color for all headers
    base_fill = PatternFill("solid", start_color="1F4E79")
    # Distinct accent colors for the three section columns
    section_fills = {
        "Progress": PatternFill("solid", start_color="1F4E79"),
        "Plans":    PatternFill("solid", start_color="375623"),
        "Problems": PatternFill("solid", start_color="833C00"),
    }
    for col, h in enumerate(HEADERS, start=1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = Font(name="Arial", bold=True, color="FFFFFF", size=11)
        c.fill = section_fills.get(h, base_fill)
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = header_border()
    ws.row_dimensions[1].height = 24


def style_data_row(ws, row_idx):
    even_fill = PatternFill("solid", start_color="EEF3F9")
    odd_fill  = PatternFill("solid", start_color="FFFFFF")
    fill = even_fill if row_idx % 2 == 0 else odd_fill

    for col in range(1, len(HEADERS) + 1):
        c = ws.cell(row=row_idx, column=col)
        c.fill = fill
        c.border = cell_border()
        c.alignment = Alignment(vertical="top", wrap_text=True)

        if col == 1:  # Date
            c.font = Font(name="Arial", size=10)
            c.alignment = Alignment(horizontal="center", vertical="top")
            c.number_format = "MM/DD/YYYY"
        elif col == 2:  # Division
            c.font = Font(name="Arial", bold=True, size=10)
        elif col == 3:  # Initiative
            c.font = Font(name="Arial", bold=True, size=10)
        else:  # Progress / Plans / Problems
            c.font = Font(name="Arial", size=10)

    ws.row_dimensions[row_idx].height = 40


def set_column_widths(ws):
    widths = {"A": 14, "B": 26, "C": 32, "D": 50, "E": 50, "F": 50}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width


def load_existing_records(ws):
    """
    Returns a set of (division, initiative) tuples already in the file.
    Used to skip exact duplicates regardless of date.
    """
    existing = set()
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        # Safely extract division (col B = index 1) and initiative (col C = index 2)
        division   = str(row[1] or "").strip() if len(row) > 1 else ""
        initiative = str(row[2] or "").strip() if len(row) > 2 else ""
        if division or initiative:
            existing.add((division, initiative))
    return existing


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 3:
        print("Usage: python docx_to_excel.py input.docx output.xlsx")
        sys.exit(1)

    in_path  = sys.argv[1]
    out_path = sys.argv[2]
    today    = date.today()

    print(f"Reading: {in_path}")
    data = parse_document(in_path)
    if not data:
        print("No entries found. Run diagnose.py to inspect document structure.")
        sys.exit(1)

    # Load or create workbook
    if os.path.exists(out_path):
        print(f"Appending to existing file: {out_path}")
        wb = load_workbook(out_path)
        ws = wb.active
        existing_records = load_existing_records(ws)
    else:
        print(f"Creating new file: {out_path}")
        wb = Workbook()
        ws = wb.active
        ws.title = "Tracker"
        for col, h in enumerate(HEADERS, start=1):
            ws.cell(row=1, column=col, value=h)
        style_header_row(ws)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = f"A1:{get_column_letter(len(HEADERS))}1"
        set_column_widths(ws)
        existing_records = set()

    # Append rows
    added = 0
    skipped = 0
    next_row = ws.max_row + 1

    for entry in data:
        key = (entry["division"], entry["initiative"])
        if key in existing_records:
            skipped += 1
            continue

        ws.cell(row=next_row, column=1, value=today)
        ws.cell(row=next_row, column=2, value=entry["division"])
        ws.cell(row=next_row, column=3, value=entry["initiative"])
        ws.cell(row=next_row, column=4, value=entry["Progress"])
        ws.cell(row=next_row, column=5, value=entry["Plans"])
        ws.cell(row=next_row, column=6, value=entry["Problems"])

        style_data_row(ws, next_row)
        existing_records.add(key)
        next_row += 1
        added += 1

    wb.save(out_path)
    print(f"✓ Done — {added} rows added, {skipped} duplicates skipped.")
    print(f"✓ Saved: {out_path}")


if __name__ == "__main__":
    main()
